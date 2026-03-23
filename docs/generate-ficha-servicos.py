"""
Gera ficha DOCX para preenchimento dos servicos do escritorio SD Legal.
Cada servico tem campos estruturados para alimentar a SonIA.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ─── Estilos ───
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ─── Capa ───
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SD LEGAL")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Ficha de Servicos — Base de Conhecimento da SonIA")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.add_run(
    "Este documento deve ser preenchido com os detalhes de cada servico "
    "prestado pelo escritorio. A informacao sera usada pela SonIA para "
    "atender os clientes pelo WhatsApp de forma precisa e profissional."
).font.size = Pt(11)

doc.add_paragraph()

instrucoes = doc.add_paragraph()
instrucoes.add_run("INSTRUCOES DE PREENCHIMENTO:").bold = True
doc.add_paragraph(
    "1. Preencha todos os campos de cada servico com linguagem clara e acessivel.",
    style="List Number",
)
doc.add_paragraph(
    "2. Nos campos de texto livre, escreva como gostaria que a SonIA explicasse ao cliente.",
    style="List Number",
)
doc.add_paragraph(
    "3. Indique valores aproximados nos honorarios (ex: 'a partir de 300 EUR').",
    style="List Number",
)
doc.add_paragraph(
    "4. Nas perguntas frequentes, inclua as duvidas reais que os clientes costumam ter.",
    style="List Number",
)
doc.add_paragraph(
    "5. Nos limites da SonIA, indique o que ela NAO deve responder e deve encaminhar ao advogado.",
    style="List Number",
)

doc.add_page_break()

# ─── Servicos ───

servicos = [
    {
        "numero": 1,
        "nome": "Autorizacao de Residencia — Primeiro Pedido",
        "codigo": "pedido_ar",
    },
    {
        "numero": 2,
        "nome": "Autorizacao de Residencia — Renovacao",
        "codigo": "renovacao_ar",
    },
    {
        "numero": 3,
        "nome": "Reagrupamento Familiar",
        "codigo": "reagrupamento_familiar",
    },
    {
        "numero": 4,
        "nome": "Nacionalidade Portuguesa",
        "codigo": "nacionalidade_pt",
    },
    {
        "numero": 5,
        "nome": "Emissao de NIF",
        "codigo": "emissao_nif",
    },
    {
        "numero": 6,
        "nome": "Constituicao de Empresa",
        "codigo": "constituicao_empresa",
    },
    {
        "numero": 7,
        "nome": "Abertura de Actividade",
        "codigo": "abertura_actividade",
    },
    {
        "numero": 8,
        "nome": "Processo Laboral",
        "codigo": "processo_laboral",
    },
    {
        "numero": 9,
        "nome": "Recurso — AR Indeferida",
        "codigo": "recurso_ar_indeferida",
    },
    {
        "numero": 10,
        "nome": "Suspensao de Saida Voluntaria",
        "codigo": "suspensao_saida_voluntaria",
    },
    {
        "numero": 11,
        "nome": "Casamento em Portugal",
        "codigo": "casamento_portugal",
    },
    {
        "numero": 12,
        "nome": "Casamento no Brasil",
        "codigo": "casamento_brasil",
    },
    {
        "numero": 13,
        "nome": "Divorcio em Portugal",
        "codigo": "divorcio_portugal",
    },
    {
        "numero": 14,
        "nome": "Divorcio no Brasil",
        "codigo": "divorcio_brasil",
    },
    {
        "numero": 15,
        "nome": "Revisao de Sentenca Estrangeira (Portugal)",
        "codigo": "revisao_sentenca_pt",
    },
    {
        "numero": 16,
        "nome": "Homologacao de Sentenca (Brasil)",
        "codigo": "homologacao_sentenca_br",
    },
    {
        "numero": 17,
        "nome": "Injuncao de Pagamento",
        "codigo": "injuncao_pagamento",
    },
    {
        "numero": 18,
        "nome": "Insolvencia de Empresa",
        "codigo": "insolvencia_empresa",
    },
    {
        "numero": 19,
        "nome": "Insolvencia Pessoal",
        "codigo": "insolvencia_pessoal",
    },
]

campos = [
    ("Descricao do servico", "Explique em 2-3 frases o que e este servico, como se estivesse a explicar a um cliente leigo."),
    ("Publico-alvo", "Quem normalmente procura este servico? (ex: imigrantes brasileiros, empresarios, etc.)"),
    ("Documentos necessarios", "Liste todos os documentos que o cliente precisa de reunir.\n1.\n2.\n3.\n4.\n5."),
    ("Etapas do processo", "Descreva passo a passo como funciona o processo.\n1.\n2.\n3.\n4.\n5."),
    ("Prazo estimado", "Quanto tempo demora tipicamente? (ex: 3-6 meses, 2-4 semanas)"),
    ("Honorarios", "Valor ou faixa de valores. (ex: 'a partir de 300 EUR', '500-800 EUR conforme complexidade')"),
    ("Forma de pagamento", "Opcoes aceites (ex: transferencia, MB Way, prestacoes, etc.)"),
    ("Perguntas frequentes dos clientes", "Quais as duvidas mais comuns que os clientes fazem sobre este servico?\n\nP1:\nR1:\n\nP2:\nR2:\n\nP3:\nR3:"),
    ("O que a SonIA PODE dizer", "Informacoes que a SonIA pode partilhar livremente com o cliente sobre este servico."),
    ("O que a SonIA NAO deve dizer", "Limites: o que deve ser encaminhado ao advogado. (ex: 'nunca dar prazos exactos de decisao do AIMA')"),
    ("Observacoes adicionais", "Notas internas, particularidades, alertas especiais sobre este servico."),
]


for servico in servicos:
    # Titulo do servico
    heading = doc.add_heading(level=1)
    run = heading.add_run(f'{servico["numero"]}. {servico["nome"]}')
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # Codigo interno
    p = doc.add_paragraph()
    p.add_run("Codigo interno: ").bold = True
    p.add_run(servico["codigo"]).font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Tabela de campos
    table = doc.add_table(rows=len(campos), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (campo_nome, campo_hint) in enumerate(campos):
        # Coluna 1: nome do campo
        cell_label = table.cell(i, 0)
        cell_label.width = Cm(5)
        p = cell_label.paragraphs[0]
        p.add_run(campo_nome).bold = True

        # Coluna 2: espaco para preencher com hint
        cell_value = table.cell(i, 1)
        cell_value.width = Cm(12)
        p = cell_value.paragraphs[0]
        run = p.add_run(campo_hint)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.size = Pt(10)

    doc.add_page_break()

# ─── Secao final: Informacoes gerais do escritorio ───

heading = doc.add_heading(level=1)
run = heading.add_run("Informacoes Gerais do Escritorio")
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

campos_gerais = [
    ("Nome completo do escritorio", ""),
    ("Morada", ""),
    ("Telefone de contacto", ""),
    ("Email geral", ""),
    ("Horario de funcionamento", "ex: Segunda a Sexta, 09:00 - 18:00"),
    ("Website", ""),
    ("Redes sociais", "Instagram, Facebook, LinkedIn, etc."),
    ("Formas de pagamento aceites", "Transferencia, MB Way, Multibanco, prestacoes, etc."),
    ("Consulta inicial", "E gratuita? Qual o valor? Presencial ou online?"),
    ("Linguas de atendimento", "ex: Portugues, Ingles, Frances, Crioulo"),
    ("Advogados do escritorio", "Nome e especialidade de cada advogado.\n\n1. Nome: _____ | Especialidade: _____\n2. Nome: _____ | Especialidade: _____"),
    ("Diferencial do escritorio", "O que distingue a SD Legal de outros escritorios? Porque devem escolher-vos?"),
    ("Tom de comunicacao desejado", "Como a SonIA deve falar com os clientes? (ex: profissional mas acolhedor, formal, descontraido)"),
    ("Frases que a SonIA NUNCA deve dizer", "Expressoes ou promessas proibidas.\n\n1.\n2.\n3."),
    ("Mensagem de boas-vindas padrao", "O que a SonIA deve dizer quando um cliente novo entra em contacto pela primeira vez?"),
]

table = doc.add_table(rows=len(campos_gerais), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

for i, (campo_nome, campo_hint) in enumerate(campos_gerais):
    cell_label = table.cell(i, 0)
    cell_label.width = Cm(5)
    p = cell_label.paragraphs[0]
    p.add_run(campo_nome).bold = True

    cell_value = table.cell(i, 1)
    cell_value.width = Cm(12)
    if campo_hint:
        p = cell_value.paragraphs[0]
        run = p.add_run(campo_hint)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.size = Pt(10)

# ─── Guardar ───
output_path = "/Users/eduardodias/escritorio-digital/docs/SD_Legal_Ficha_Servicos_SonIA.docx"
doc.save(output_path)
print(f"Ficheiro gerado: {output_path}")
